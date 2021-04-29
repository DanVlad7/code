from docx import Document
from docx.shared import Inches
import pyttsx3


document = Document()

#functions

def speak(text):
    pyttsx3.speak(text)

def workexperience():
    p = document.add_paragraph()
    company = input('Enter company ')
    from_date = input('From date')
    to_date = input('To date ')
    p.add_run(company + ' ').bold = True
    p.add_run(from_date + ' - ' + to_date + '\n').italic = True
    experience_details = input(
        'Describe your experience at ' + company)
    p.add_run(experience_details)

def skills():
    p = document.add_paragraph()
    p.style = 'List Bullet'
    skill = input('Tell me your skill')
    p.add_run(skill)

#[picture]
document.add_picture(
    'dan.jpg',
    width = Inches(2.0)
)

#details
name = input("What is your name? ")
phone_number = input("What is your phone number]? ")
email = input("What is your email? ")

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

#about
document.add_heading('About me')
about_me =input('Tell me about yourself? ')

document.add_paragraph(about_me)

#work experience
document.add_heading('Work experience')
workexperience()

#more experiences
while True:
    has_more_experiences = input(
        'Do you have other experiences?')
    if has_more_experiences.lower() == 'yes':
        workexperience()
    elif has_more_experiences.lower() == 'no':
        break

#skills
document.add_heading('Skills')
skills()

while True:
    has_more_skills = input(
        'Do you have more skills?')
    if has_more_skills.lower() == 'yes':
        skills()
    elif has_more_skills.lower() == 'no':
        break

#footer
section = document.sections[0]
footer = section.footer
p = footer.add_paragraphs[0]
p.text = 'CV generated automatically'

document.save('cv.docx')

